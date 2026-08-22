#!/usr/bin/env python3
"""
真实 Word 修订模式（Track Changes）写入器：无批注、无摘要。

用途：当用户要“类似 Word 审阅模式直接改出来”“不要批注”时使用。
它写入 OOXML 的 <w:ins>/<w:del>，并在 settings.xml 开启 <w:trackRevisions/>。

输入 JSON：
[
  {"paragraph": 5, "new_text": "替换后的段落文本"},
  {"paragraph": 26, "delete": true}
]

输出：真实修订标记 docx。Word 打开后可在“审阅”中接受/拒绝。
"""
import argparse
import json
import re
import zipfile
import tempfile
from pathlib import Path
from datetime import datetime, timezone
import difflib

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

AUTHOR = "Hermes AI"


def tokenize_cn(s: str):
    """中文短语级 token，避免整段删除+整段插入。"""
    return re.findall(r"[\u4e00-\u9fff]+|[A-Za-z0-9]+(?:\.[0-9]+)?|[“”‘’《》（）()、，。；：！？,.!?;:\-—…]+|\s+", s or "")


def clear_content(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def rpr(font="宋体", size=24):
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size)); rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(size)); rPr.append(szCs)
    return rPr


def run_text(text, deleted=False, font="宋体"):
    r = OxmlElement("w:r")
    r.append(rpr(font))
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def make_change(tag, text, rev_id, date, font="宋体"):
    e = OxmlElement(f"w:{tag}")
    e.set(qn("w:id"), str(rev_id))
    e.set(qn("w:author"), AUTHOR)
    e.set(qn("w:date"), date)
    e.append(run_text(text, deleted=(tag == "del"), font=font))
    return e


def apply_diff(paragraph, new_text, rev_counter, date, font="宋体"):
    old_text = paragraph.text
    clear_content(paragraph)
    old_t = tokenize_cn(old_text)
    new_t = tokenize_cn(new_text)
    sm = difflib.SequenceMatcher(a=old_t, b=new_t, autojunk=False)
    p = paragraph._p
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        old = "".join(old_t[i1:i2])
        new = "".join(new_t[j1:j2])
        if tag == "equal":
            if old:
                p.append(run_text(old, deleted=False, font=font))
        elif tag == "delete":
            if old:
                p.append(make_change("del", old, next(rev_counter), date, font))
        elif tag == "insert":
            if new:
                p.append(make_change("ins", new, next(rev_counter), date, font))
        elif tag == "replace":
            if old:
                p.append(make_change("del", old, next(rev_counter), date, font))
            if new:
                p.append(make_change("ins", new, next(rev_counter), date, font))


def apply_delete(paragraph, rev_counter, date, font="宋体"):
    old_text = paragraph.text
    clear_content(paragraph)
    if old_text:
        paragraph._p.append(make_change("del", old_text, next(rev_counter), date, font))


def enable_track_revisions(docx_path: Path):
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(td)
        settings = td / "word" / "settings.xml"
        xml = settings.read_text(encoding="utf-8")
        insert = ""
        if "<w:trackRevisions" not in xml:
            insert += "<w:trackRevisions/>"
        if "<w:revisionView" not in xml:
            insert += '<w:revisionView w:markup="1" w:comments="1" w:insDel="1" w:formatting="1"/>'
        if insert:
            xml = xml.replace("</w:settings>", insert + "</w:settings>")
            settings.write_text(xml, encoding="utf-8")
        tmp = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for f in td.rglob("*"):
                if f.is_file():
                    zout.write(f, f.relative_to(td).as_posix())
        tmp.replace(docx_path)


def verify(docx_path: Path):
    with zipfile.ZipFile(docx_path) as z:
        names = z.namelist()
        document_xml = z.read("word/document.xml").decode("utf-8")
        settings_xml = z.read("word/settings.xml").decode("utf-8")
    return {
        "w_ins": document_xml.count("<w:ins "),
        "w_del": document_xml.count("<w:del "),
        "track_revisions": "<w:trackRevisions" in settings_xml,
        "revision_view": "<w:revisionView" in settings_xml,
        "has_comments_xml": any("comments.xml" in n for n in names),
    }


def counter(start=1):
    n = start
    while True:
        yield n
        n += 1


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("revisions_json")
    ap.add_argument("-o", "--output")
    args = ap.parse_args()

    src = Path(args.docx)
    out = Path(args.output) if args.output else src.with_name(src.stem + "_修订无批注" + src.suffix)
    revisions = json.loads(Path(args.revisions_json).read_text(encoding="utf-8"))

    doc = Document(src)
    date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_counter = counter(1)

    for rev in revisions:
        idx = int(rev["paragraph"]) - 1
        if idx < 0 or idx >= len(doc.paragraphs):
            raise IndexError(f"paragraph {idx+1} out of range")
        if rev.get("delete"):
            apply_delete(doc.paragraphs[idx], rev_counter, date)
        else:
            apply_diff(doc.paragraphs[idx], rev["new_text"], rev_counter, date)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    enable_track_revisions(out)
    info = verify(out)
    print(json.dumps({"output": str(out), **info}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
