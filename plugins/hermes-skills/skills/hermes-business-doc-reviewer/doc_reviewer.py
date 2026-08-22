#!/usr/bin/env python3
"""
文档审阅器 — 直接在 Word 原文件里加修订标记和批注。
用删除线+高亮标记原文，段落末加批注标签，文末汇总所有审阅意见。
"""
import sys, os, json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH


def strike_run(para, original_text):
    """给段落中匹配的文字加红色删除线"""
    for run in para.runs:
        if original_text in run.text:
            idx = run.text.find(original_text)
            end = idx + len(original_text)
            
            # Split into three parts using XML manipulation
            before_text = run.text[:idx]
            struck_text = run.text[idx:end]
            after_text = run.text[end:]
            
            # Clear current run and set to "before" text
            run.text = before_text if before_text else ''
            
            # Insert struck run after current
            struck_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            strike = OxmlElement('w:strike')
            strike.set(qn('w:val'), 'true')
            rPr.append(strike)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), 'CC0000')
            rPr.append(color)
            struck_el.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = struck_text
            struck_el.append(t)
            
            run._element.addnext(struck_el)
            
            # Insert after text if any
            if after_text:
                after_el = OxmlElement('w:r')
                t2 = OxmlElement('w:t')
                t2.set(qn('xml:space'), 'preserve')
                t2.text = after_text
                after_el.append(t2)
                struck_el.addnext(after_el)
            
            return True
    return False


def highlight_text(paragraph, original_text):
    """高亮段落中的指定文字"""
    for run in paragraph.runs:
        if original_text in run.text:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return True
    return False


def add_inline_comment(paragraph, comment_text, label="📝"):
    """在段落末尾添加批注标签"""
    run = paragraph.add_run(' ')
    run.font.size = Pt(8)
    run = paragraph.add_run(f'  {label}批注: {comment_text}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True


def apply_review(doc_path, reviews, output_path=None):
    """应用审阅意见到文档"""
    if output_path is None:
        base, ext = os.path.splitext(doc_path)
        output_path = f"{base}_审阅{ext}"
    
    doc = Document(doc_path)
    
    # 统计
    summary_lines = []
    severe = moderate = suggestion = 0
    
    for rev in reviews:
        para_idx = rev.get('paragraph', 0) - 1
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            print(f"⚠️ 段落 {para_idx+1} 不存在，跳过")
            continue
        
        para = doc.paragraphs[para_idx]
        rev_type = rev.get('type', 'comment')
        orig = rev.get('original_text', '')
        text = rev.get('text', '')
        suggestion_text = rev.get('suggestion', '')
        
        label_map = {'delete': '🔴', 'modify': '🟡', 'highlight': '🟡', 'comment': '🟢'}
        label = label_map.get(rev_type, '📝')
        level = {'delete': 2, 'modify': 1, 'highlight': 1, 'comment': 0}
        sev = level.get(rev_type, 0)
        if sev == 2: severe += 1
        elif sev == 1: moderate += 1
        else: suggestion += 1
        
        # 应用标记
        if rev_type == 'delete' and orig:
            strike_run(para, orig)
        
        elif rev_type == 'modify' and orig:
            strike_run(para, orig)
            full_comment = f"建议改为「{suggestion_text}」"
            if text:
                full_comment += f"（{text}）"
            add_inline_comment(para, full_comment, label)
        
        elif rev_type == 'highlight' and orig:
            highlight_text(para, orig)
            add_inline_comment(para, text, label)
        
        elif rev_type == 'comment':
            add_inline_comment(para, text, label)
        
        # 添加到摘要
        pnum = f"第{para_idx+1}段"
        snippet = para.text[:60].strip() if para.text else "(空段落)"
        summary_lines.append(f"{label} {pnum}：{text[:80]}")
    
    # 文末添加审阅摘要
    if summary_lines:
        doc.add_paragraph('')
        doc.add_paragraph('─' * 60)
        heading = doc.add_heading('审阅摘要', level=2)
        
        stats = doc.add_paragraph()
        stats.add_run(f'🔴 严重 {severe} 处  ').bold = True
        stats.add_run(f'🟡 注意 {moderate} 处  ').bold = True
        stats.add_run(f'🟢 建议 {suggestion} 处').bold = True
        
        for line in summary_lines:
            p = doc.add_paragraph(line)
            p.style = doc.styles['List Bullet'] if 'List Bullet' in [s.name for s in doc.styles] else doc.styles['Normal']
        
        doc.add_paragraph(f'\n审阅时间：{__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")}')
    
    doc.save(output_path)
    print(f"✅ 审阅完成：{output_path}")
    print(f"   🔴{severe} 🟡{moderate} 🟢{suggestion} 共 {len(reviews)} 条")
    return output_path


def extract_full_text(doc_path):
    """提取文档全文，带段落编号"""
    doc = Document(doc_path)
    lines = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            lines.append(f"[P{i+1}] {para.text}")
    return '\n'.join(lines)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python3 doc_reviewer.py <docx路径> [审阅JSON路径]")
        print("  或：python3 doc_reviewer.py <docx路径> --extract")
        sys.exit(1)

    if '--extract' in sys.argv:
        print(extract_full_text(sys.argv[1]))
    else:
        doc_path = sys.argv[1]
        reviews = json.loads(sys.stdin.read()) if len(sys.argv) < 3 else json.load(open(sys.argv[2]))
        apply_review(doc_path, reviews)
